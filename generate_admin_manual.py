from docx import Document
from docx.shared import Pt, Inches

def generate_admin_manual():
    doc = Document()
    
    doc.add_heading('Buku Panduan Penggunaan Website Warkop Ayah - Khusus Admin', 0)
    
    doc.add_paragraph('Dokumen ini berisi panduan dan langkah-langkah penggunaan website Warkop Ayah yang dikhususkan untuk Administrator.')
    
    doc.add_heading('1. Cara Login Admin', level=1)
    doc.add_paragraph('Untuk masuk ke halaman dashboard admin, ikuti langkah berikut:')
    doc.add_paragraph('a. Pastikan server aplikasi (backend dan frontend) sudah berjalan (klik file start_restoran.bat).')
    doc.add_paragraph('b. Buka browser dan kunjungi link: http://127.0.0.1:8191/admin/login.html')
    doc.add_paragraph('c. Masukkan Username dan Password Anda.')
    doc.add_paragraph('   - Username default: admin')
    doc.add_paragraph('   - Password default: admin123')
    doc.add_paragraph('d. Klik tombol Login. Jika berhasil, Anda akan langsung diarahkan ke halaman Dashboard Admin.')

    doc.add_heading('2. Manajemen Menu (Menambah Menu Baru)', level=1)
    doc.add_paragraph('Untuk menambahkan menu baru ke dalam sistem, ikuti langkah berikut:')
    doc.add_paragraph('a. Pada menu navigasi di sebelah kiri, klik "Manajemen Menu".')
    doc.add_paragraph('b. Klik tombol "+ Tambah Menu".')
    doc.add_paragraph('c. Isi form yang disediakan:')
    doc.add_paragraph('   - Nama: Masukkan nama menu (contoh: Kopi Hitam).')
    doc.add_paragraph('   - Harga: Masukkan harga dalam angka tanpa titik (contoh: 5000).')
    doc.add_paragraph('   - Stok: Masukkan jumlah stok awal.')
    doc.add_paragraph('   - Jenis/Kategori: Pilih kategori yang sesuai (makanan, minuman, dll).')
    doc.add_paragraph('   - Foto: Pilih foto/gambar yang sudah di-upload sebelumnya pada Manajemen Gambar.')
    doc.add_paragraph('d. Setelah semua data terisi dengan benar, klik "Simpan Menu". Menu baru akan langsung tampil di halaman pelanggan.')

    doc.add_heading('3. Manajemen Menu (Mengubah & Menghapus Menu)', level=1)
    doc.add_paragraph('a. Mengubah Menu (Edit): Pada halaman "Manajemen Menu", cari menu yang ingin diubah, lalu klik tombol Edit (ikon pensil kuning) di baris menu tersebut. Ubah data yang diperlukan, lalu simpan.')
    doc.add_paragraph('b. Menghapus Menu: Pada halaman "Manajemen Menu", cari menu yang ingin dihapus, lalu klik tombol Hapus (ikon tempat sampah merah) di baris menu tersebut. Konfirmasi penghapusan saat diminta.')

    doc.add_heading('4. Manajemen Gambar (Mengunggah Foto Menu)', level=1)
    doc.add_paragraph('Sebelum Anda bisa menggunakan gambar saat menambah menu baru, gambar harus di-upload terlebih dahulu.')
    doc.add_paragraph('a. Klik "Manajemen Gambar" pada menu navigasi sebelah kiri.')
    doc.add_paragraph('b. Klik tombol Pilih File (Browse) dan pilih foto dari komputer Anda.')
    doc.add_paragraph('c. Klik Upload. Gambar yang sukses diunggah akan muncul di daftar gambar di bawahnya, dan siap dipilih saat menambah/edit menu.')

    doc.add_heading('5. Manajemen Kategori', level=1)
    doc.add_paragraph('a. Klik "Manajemen Kategori" pada navigasi.')
    doc.add_paragraph('b. Anda dapat menambahkan kategori baru yang nantinya bisa digunakan saat membuat menu baru, atau menghapus/mengubah kategori yang sudah ada.')

    doc.add_paragraph('\n-- Panduan ini dibuat khusus untuk mempermudah operasional Admin Warkop Ayah --').italic = True

    doc.save('Panduan_Admin_Warkop_Ayah.docx')
    print('Dokumen Panduan Admin berhasil dibuat!')

if __name__ == "__main__":
    generate_admin_manual()
