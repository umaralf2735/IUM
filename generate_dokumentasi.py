from docx import Document
from docx.shared import Pt, Inches

def generate_dokumentasi():
    doc = Document()
    
    # Judul Dokumen
    doc.add_heading('Dokumentasi Program Frontend & Backend Warkop Ayah', 0)
    
    doc.add_paragraph('Dokumen ini berisi penjelasan teknis mengenai arsitektur perangkat lunak, struktur frontend dan backend, serta pembagian tugas (Person In Charge - PIC) dalam pengembangan aplikasi Warkop Ayah.')
    
    # 1. PIC
    doc.add_heading('1. Person In Charge (PIC)', level=1)
    doc.add_paragraph('Bagian ini menjelaskan pembagian tugas / peran masing-masing anggota dalam pengembangan sistem Warkop Ayah:')
    
    # Table for PIC
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nama Anggota'
    hdr_cells[1].text = 'NIM'
    hdr_cells[2].text = 'Bagian (Frontend / Backend / Fullstack / Database)'
    
    # Template baris kosong untuk diisi user
    for i in range(3):
        row_cells = table.add_row().cells
        row_cells[0].text = f'[Nama Anggota ke-{i+1}]'
        row_cells[1].text = '[NIM]'
        row_cells[2].text = '[Tugas]'

    # 2. Arsitektur Frontend
    doc.add_heading('2. Dokumentasi Frontend (HTML/CSS/JS)', level=1)
    doc.add_paragraph('Frontend dibangun menggunakan Vanilla HTML, CSS, dan JavaScript murni tanpa framework tambahan agar lebih ringan dan cepat. Struktur pada frontend:')
    
    doc.add_heading('A. Halaman Kasir / Pelanggan', level=2)
    doc.add_paragraph('- index.html : Layar utama untuk menampilkan daftar menu, kategori makanan, dan fungsi keranjang belanja.\n- css/style.css : Pengaturan styling (warna, layout, desain responsif) untuk sisi pelanggan.\n- js/main.js : Berisi logika JavaScript untuk mengambil (fetch) layanan data lewat API Backend, menampilkan ke layar, memfilter menu makanan, dan memproses perhitungan pembayaran dan kembalian keranjang.')

    doc.add_heading('B. Halaman Admin Dashboard', level=2)
    doc.add_paragraph('- admin/login.html : Layar antarmuka untuk masuk ke dalam dasbor. Dilengkapi sistem keamanan validasi kredensial.\n- admin/index.html : Layar dashboard bagi admin untuk memanajemen data Menu, Kategori, dan Daftar Gambar Etalase.\n- admin/css/admin.css : Konfigurasi gaya tampilan spesifik untuk Dashboard Admin dan Sidebar.\n- admin/js/dashboard.js : Logika manajemen State Dashboard, berisi CRUD (Create, Read, Update, Delete) yang menembak langsung backend API.')

    # 3. Arsitektur Backend
    doc.add_heading('3. Dokumentasi Backend (Python Flask API)', level=1)
    doc.add_paragraph('Backend dikembangkan dengan bahasa Python menggunakan micro-framework Flask. Backend ini murni bertindak sebagai RESTful API yang melayani request JSON dari Frontend.')
    
    doc.add_heading('Struktur Endpoint Routes', level=2)
    p = doc.add_paragraph()
    p.add_run('- Authentication : ').bold = True
    p.add_run('POST /api/login untuk mencocokkan kredensial default admin.')
    
    p2 = doc.add_paragraph()
    p2.add_run('- Products (Menu) : ').bold = True
    p2.add_run('GET /api/products untuk memuat data menu, POST /api/products untuk menambah, PUT /api/products/<id> untuk mengubah, dan DELETE /api/products/<id> untuk menghapus makanan dari menu.')

    p3 = doc.add_paragraph()
    p3.add_run('- Categories : ').bold = True
    p3.add_run('GET /api/categories, POST /api/categories untuk manajemen kategori unik (contoh: Makanan, Minuman).')

    p4 = doc.add_paragraph()
    p4.add_run('- Images Upload : ').bold = True
    p4.add_run('POST /api/upload untuk mengungkit gambar baru dari komputer admin langsung ke server, untuk kemudian diassign ke produk.')

    p5 = doc.add_paragraph()
    p5.add_run('- Checkout (Orders) : ').bold = True
    p5.add_run('POST /api/orders untuk merekam transaksi yang sudah selesai dibayar pada sisi kasir, tersambung ke order history database.')

    # 4. Database
    doc.add_heading('4. Dokumentasi Database MySQL', level=1)
    doc.add_paragraph('Database berelasi ter-struktur menggunakan MySQL. Python menggunakan teknologi SQLAlchemy (ORM) / mysql-connector untuk menghubungkan diri. Struktur tabel berisi:\n\n- Tabel users (id, username, password_hash): Menyimpan hak akses para admin.\n- Tabel categories (id, name): Master data pengelompokan menu.\n- Tabel products (id, name, price, category_id, image_url, etc): Data makanan dan minuman. Berelasi Foreign Key ke -> tabel categories.\n- Tabel images (id, url, filename): Entitas daftar gambar yang tersedia pada galeri internal web.\n- Tabel orders dan order_items: Tempat riwayat rekam jejak pesanan.')

    # Penutup
    doc.add_paragraph('\n-- Dokumen teknis digenerate secara otomatis untuk mempermudah penilaian. Silakan melengkapi nama Anggota/PIC pada tabel di atas. --').italic = True

    doc.save('Dokumentasi_Program_PIC.docx')
    print('Dokumentasi_Program_PIC.docx sukses dibuat!')

if __name__ == "__main__":
    generate_dokumentasi()
