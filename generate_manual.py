from docx import Document
from docx.shared import Pt, Inches

def generate_manual():
    doc = Document()
    
    # Title
    doc.add_heading('Buku Panduan Penggunaan Website Warkop Ayah', 0)
    
    doc.add_paragraph('Dokumen ini berisi langkah-langkah detail penggunaan fitur website Warkop Ayah baik dari kacamata pelanggan maupun dari sisi admin.')
    
    doc.add_heading('A. Akses Website', level=1)
    doc.add_paragraph('1. Buka file start_restoran.bat untuk menyalakan server secara otomatis.\n2. Tampilan Kasir/Pelanggan dapat diakses di browser melalui link: http://127.0.0.1:8191\n3. Tampilan Admin Dashboard dapat diakses di link: http://127.0.0.1:8191/admin/login.html')

    doc.add_heading('B. Panduan Penggunaan Halaman Pelanggan (Kasir)', level=1)
    
    doc.add_heading('1. Menambahkan Pesanan ke Keranjang', level=2)
    p = doc.add_paragraph()
    p.add_run('Pada bagian ')
    p.add_run('MENU').bold = True
    p.add_run(', pilih menu yang ingin dibeli, kemudian klik tombol ')
    p.add_run('BELI SEKARANG').bold = True
    p.add_run('. Menu otomatis akan masuk ke keranjang belanja Anda. Anda dapat menyaring menu secara spesifik berdasarkan kategori Makanan, Minuman, atau Snack menggunakan urutan tombol filter di atas daftar Menu.')
    
    doc.add_heading('2. Mengecek & Mengubah Keranjang (Cart)', level=2)
    doc.add_paragraph('Akan muncul notifikasi pada ikon Keranjang Kuning (melayang) di sudut kanan bawah setiap kali menu ditambah. Klik tombol tersebut untuk membuka jendela menu Keranjang. Di dalam keranjang, Anda bisa mengubah jumlah satuan (Quantity) secara leluasa atau menghapus produk dari pesanan.')
    
    doc.add_heading('3. Melakukan Pembayaran', level=2)
    doc.add_paragraph('Di tab Keranjang yang sudah Anda buka:\n1. Klik tombol "Proses Pembayaran".\n2. Cek Total Tagihan.\n3. Masukkan jumlah uang tunai yang diberikan langsung ke kasir.\n4. Sistem akan secara otomatis menghitung kembalian uang jika uang berlebih.\n5. Jika cukup, pastikan Anda menekan tombol "Konfirmasi & Bayar" dan pesanan berhasil terekam!')

    doc.add_heading('C. Panduan Penggunaan Dashboard Admin', level=1)
    
    doc.add_heading('1. Login Admin', level=2)
    doc.add_paragraph('Masuk ke Dashboard via link admin (http://127.0.0.1.../admin/login.html).\nIsikan Username (default: admin) dan Password (default: admin123).\nSetelah login berhasil, Anda akan dialihkan ke layar utama admin yang berisi manajemen warung.')

    doc.add_heading('2. Mengelola Menu & Produk', level=2)
    doc.add_paragraph('Pilih menu "Manajemen Menu" di panel kiri.\n- Tambah Menu Baru: Klik tombol hijau "+ Tambah Menu". Isi form yang diminta seperti Nama, Harga, Stok, Jenis, dan Foto dari galeri internal. Klik "Simpan Menu".\n- Edit Menu: Klik ikon Pensil Kuning di urutan makanan.\n- Hapus Menu: Klik ikon Trash/Sampah warna merah.')

    doc.add_heading('3. Menambah Gambar Produk untuk Etalase', level=2)
    doc.add_paragraph('Anda juga diberi wewenang khusus untuk dapat mengunggah (upload) gambar-gambar ke penyimpanan internet lokal tanpa perlu akses file proyek code.\n1. Buka tabel "Manajemen Gambar".\n2. Klik Browse untuk mencari foto, kemudian klik Upload.\n3. Gambar sukses di-unggah dan akan tersedia pada form dropdown Add/Edit Menu!')
    
    doc.add_heading('4. Mengelola Kategori', level=2)
    doc.add_paragraph('Fungsi ini memungkinkan Admin menciptakan variasi kategori selain standard yang tersedia. Ini bisa diakses pada panel "Manajemen Kategori".')

    doc.add_heading('D. Fitur Opsional Maps', level=1)
    doc.add_paragraph('Anda selalu bisa meng-klik tombol navigasi LOKASI KAMI bagian atas untuk mendapatkan kontak & maps ke alamat restoran secara real-time dan akurat.')
    
    p2 = doc.add_paragraph()
    p2.add_run('\n-- Terimakasih, silakan simpan panduan di atas --').italic = True

    doc.save('Buku_Manual_Warkop_Ayah.docx')
    print('Docx generated!')

if __name__ == "__main__":
    generate_manual()
